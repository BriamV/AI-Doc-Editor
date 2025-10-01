"""
Text extraction service for document processing
T-04-ST2: Extract text from PDF, DOCX, and Markdown files for RAG pipeline
"""

import logging
from pathlib import Path

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None  # type: ignore

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Custom exception for text extraction failures"""

    pass


class TextExtractor:
    """
    Service for extracting text from various document formats
    Supports PDF, DOCX, and Markdown files
    """

    def __init__(self):
        """Initialize text extractor with format validation"""
        self._validate_dependencies()

    def _validate_dependencies(self) -> None:
        """Validate that required libraries are installed"""
        if PdfReader is None:
            logger.warning("pypdf not installed - PDF extraction will not work")
        if DocxDocument is None:
            logger.warning("python-docx not installed - DOCX extraction will not work")

    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using pypdf

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content

        Raises:
            TextExtractionError: If extraction fails
        """
        if PdfReader is None:
            raise TextExtractionError("pypdf library not installed")

        try:
            logger.info(f"Starting PDF extraction: {file_path}")

            if not Path(file_path).exists():
                raise TextExtractionError(f"File not found: {file_path}")

            # Read PDF file
            reader = PdfReader(file_path)

            # Extract text from all pages
            text_parts = []
            page_count = len(reader.pages)

            logger.debug(f"Processing {page_count} pages from PDF")

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    logger.debug(f"Extracted {len(page_text)} chars from page {page_num}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue

            if not text_parts:
                raise TextExtractionError("No text could be extracted from PDF")

            extracted_text = "\n\n".join(text_parts)
            logger.info(f"PDF extraction complete: {len(extracted_text)} characters")

            return extracted_text

        except TextExtractionError:
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")

    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content

        Raises:
            TextExtractionError: If extraction fails
        """
        if DocxDocument is None:
            raise TextExtractionError("python-docx library not installed")

        try:
            logger.info(f"Starting DOCX extraction: {file_path}")

            if not Path(file_path).exists():
                raise TextExtractionError(f"File not found: {file_path}")

            # Read DOCX file
            doc = DocxDocument(file_path)

            # Extract text from all paragraphs
            text_parts = []
            paragraph_count = len(doc.paragraphs)

            logger.debug(f"Processing {paragraph_count} paragraphs from DOCX")

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            if not text_parts:
                raise TextExtractionError("No text could be extracted from DOCX")

            extracted_text = "\n\n".join(text_parts)
            logger.info(f"DOCX extraction complete: {len(extracted_text)} characters")

            return extracted_text

        except TextExtractionError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")

    def extract_from_markdown(self, file_path: str) -> str:
        """
        Extract text from Markdown file (simple file read with UTF-8 encoding)

        Args:
            file_path: Path to Markdown file

        Returns:
            Extracted text content

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            logger.info(f"Starting Markdown extraction: {file_path}")

            if not Path(file_path).exists():
                raise TextExtractionError(f"File not found: {file_path}")

            # Read markdown file with UTF-8 encoding
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            if not text.strip():
                raise TextExtractionError("Markdown file is empty")

            logger.info(f"Markdown extraction complete: {len(text)} characters")

            return text

        except TextExtractionError:
            raise
        except UnicodeDecodeError as e:
            logger.error(f"Markdown encoding error: {str(e)}")
            raise TextExtractionError(f"Failed to decode Markdown file (encoding issue): {str(e)}")
        except Exception as e:
            logger.error(f"Markdown extraction failed: {str(e)}")
            raise TextExtractionError(f"Failed to extract text from Markdown: {str(e)}")

    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Main extraction method that dispatches to specific extractors

        Args:
            file_path: Path to file
            file_type: Type of file (pdf, docx, md)

        Returns:
            Extracted text content

        Raises:
            TextExtractionError: If extraction fails or file type is unsupported
        """
        file_type_lower = file_type.lower()

        logger.info(f"Extracting text from {file_type_lower} file: {file_path}")

        # Dispatch to appropriate extractor
        if file_type_lower == "pdf":
            return self.extract_from_pdf(file_path)
        elif file_type_lower == "docx":
            return self.extract_from_docx(file_path)
        elif file_type_lower in ["md", "markdown"]:
            return self.extract_from_markdown(file_path)
        else:
            raise TextExtractionError(f"Unsupported file type: {file_type}")

    def validate_text_quality(self, text: str, min_length: int = 10) -> bool:
        """
        Validate that extracted text meets quality requirements

        Args:
            text: Extracted text
            min_length: Minimum acceptable text length

        Returns:
            True if text is valid, False otherwise
        """
        if not text or not text.strip():
            logger.warning("Extracted text is empty")
            return False

        if len(text.strip()) < min_length:
            logger.warning(f"Extracted text too short: {len(text)} chars (min: {min_length})")
            return False

        return True
